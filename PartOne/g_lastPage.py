from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches  # Import the Inches object
from . import c_queryParser
from docx.shared import Pt
import os


def createHeaderAndFooter(docNumber):
    headerList, footerList = c_queryParser.headerFooterParsing(docNumber)
    document = Document('middle.docx')
    style = document.styles['Normal']
    font = style.font
    font.name = 'Cascadia Code'
    font.complex_script = True
    font.rtl = True
    font.size = Pt(5)
    if len(document.tables) > 2:
        table = document.tables[2]
        for i in range(len(footerList)):
            for j in range(len(table.columns)):
                cell = table.cell(i, j)
                cell.text = str(footerList[i])
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(document.tables) > 0:
        header_table = document.tables[0]
    if len(header_table.rows) * len(header_table.columns) == len(headerList):
        for j, data in enumerate(headerList):
            row_index = j // len(header_table.columns)
            col_index = j % len(header_table.columns)
            header_cell = header_table.cell(row_index, col_index)
            header_cell.text = data
            for paragraph in header_cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        print("number of cells does not match the length of headerList.")
    document.save('lastPage.docx')


def createLastTable(docNumber, lastList, imgName):
    finalDataList = lastList
    # document = Document('pt3.docx')
    document = Document(os.path.join(os.path.dirname(__file__), 'pt3-QR.docx'))
    style = document.styles['Normal']
    font = style.font
    font.name = 'Cascadia Code'
    font.complex_script = True
    font.rtl = True
    font.size = Pt(5)
    if len(document.tables) > 1:
        table = document.tables[1]
        num_rows = len(table.rows)
        if num_rows >= len(finalDataList):
            for j in range(len(finalDataList)):
                for k in range(len(finalDataList[j])):
                    cell = table.cell(j, k)
                    cell.margin_top = Pt(0)
                    cell.margin_right = Pt(0)
                    cell.text = finalDataList[j][k]
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # # $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$
    # if len(document.tables) >= 4:
    #     table = document.tables[3]
    #     cell = table.cell(0, 0)
    #     for paragraph in cell.paragraphs:
    #         for run in paragraph.runs:
    #             run.clear()
    #         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     img_path = "C:\\Users\\BAB AL SAFA\\Desktop\\Organized-TM-genrator\\PartOne\\code.png"
    #     cell.paragraphs[0].clear()
    #     run = cell.paragraphs[0].add_run()
    #     run.add_picture(img_path)
    # else:
    #     print("There are not enough tables in the document to add an image to the fourth table.")
    # # $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$
    # Check if there are at least 4 tables in the document
    if len(document.tables) >= 4:
        table = document.tables[3]
        cell = table.cell(0, 0)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.clear()

        img_path = "C:\\Users\\BAB AL SAFA\\Desktop\\Organized-TM-genrator\\PartOne\\" + imgName
        cell.paragraphs[0].clear()

        # Set the cell width to match the image width
        cell.width = Inches(1.2)  # Adjust the width as needed

        run = cell.paragraphs[0].add_run()
        run.add_picture(img_path)

        # Set paragraph alignment to center
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        print("There are not enough tables in the document to add an image to the fourth table.")

    document.save('middle.docx')
    createHeaderAndFooter(docNumber)


def createLastTableEmpty(docNumber, imgName):
    headerList, footerList = c_queryParser.headerFooterParsing(docNumber)
    # document = Document('pt3.docx')
    document = Document(os.path.join(os.path.dirname(__file__), 'pt3-QR.docx'))
    style = document.styles['Normal']
    font = style.font
    font.name = 'Cascadia Code'
    font.complex_script = True
    font.rtl = True
    font.size = Pt(5)
    if len(document.tables) > 2:
        table = document.tables[2]
        for i in range(len(footerList)):
            for j in range(len(table.columns)):
                cell = table.cell(i, j)
                cell.text = str(footerList[i])
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(document.tables) > 0:
        header_table = document.tables[0]
    if len(header_table.rows) * len(header_table.columns) == len(headerList):
        for j, data in enumerate(headerList):
            row_index = j // len(header_table.columns)
            col_index = j % len(header_table.columns)
            header_cell = header_table.cell(row_index, col_index)
            header_cell.text = data
            for paragraph in header_cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        print("number of cells does not match the length of headerList.")

    # # $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$
    # if len(document.tables) >= 4:
    #     table = document.tables[3]
    #     cell = table.cell(0, 0)
    #     for paragraph in cell.paragraphs:
    #         for run in paragraph.runs:
    #             run.clear()
    #         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     img_path = "C:\\Users\\BAB AL SAFA\\Desktop\\Organized-TM-genrator\\PartOne\\code.png"
    #     cell.paragraphs[0].clear()
    #     run = cell.paragraphs[0].add_run()
    #     run.add_picture(img_path)
    # else:
    #     print("There are not enough tables in the document to add an image to the fourth table.")
    # # $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$

    # Check if there are at least 4 tables in the document
    if len(document.tables) >= 4:
        table = document.tables[3]
        cell = table.cell(0, 0)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.clear()

        img_path = "C:\\Users\\BAB AL SAFA\\Desktop\\Organized-TM-genrator\\PartOne\\"+imgName
        cell.paragraphs[0].clear()

        # Set the cell width to match the image width
        cell.width = Inches(1.2)  # Adjust the width as needed

        run = cell.paragraphs[0].add_run()
        run.add_picture(img_path)

        # Set paragraph alignment to center
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        print("There are not enough tables in the document to add an image to the fourth table.")

    document.save('lastPage.docx')

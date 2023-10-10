from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from . import c_queryParser, g_lastTable
import os 

def createMainsPages(docNumber):
    finalDataList = c_queryParser.rowsParsing(docNumber)
    headerList, footerList = c_queryParser.headerFooterParsing(docNumber)
    chunk_size = 10
    data_chunks = [finalDataList[i:i + chunk_size]
                   for i in range(0, len(finalDataList), chunk_size)]
    file_names = []

    for i, chunk in enumerate(data_chunks):
        # document = Document('pt1.docx')
        document = Document(os.path.join(os.path.dirname(__file__), 'pt1.docx'))
        style = document.styles['Normal']
        font = style.font
        font.name = 'Cascadia Code'
        font.complex_script = True
        font.rtl = True
        font.size = Pt(5)
        if len(chunk) < 7 and (i == len(data_chunks) - 1):
            g_lastTable.createLastTable(docNumber, chunk)
            break
        elif len(chunk) == 10 and (i == len(data_chunks) - 1):
            g_lastTable.createLastTableEmpty(docNumber)
            # Dont Break and Let It Continue Creating last Chunck
        elif (len(chunk) >= 7 and len(chunk) < 11) and (i == len(data_chunks) - 1):
            if len(document.tables) > 1:
                table = document.tables[1]
                num_rows = len(table.rows)
                if num_rows >= len(chunk):
                    for j in range(len(chunk)):
                        for k in range(len(chunk[j])):
                            cell = table.cell(j, k)
                            cell.margin_top = Pt(0)
                            cell.margin_right = Pt(0)
                            cell.text = chunk[j][k]
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if len(document.tables) > 0:
                header_table = document.tables[0]
                if len(header_table.rows) * len(header_table.columns) == len(headerList):
                    for j, data in enumerate(headerList):
                        row_index = j // len(header_table.columns)
                        col_index = j % len(header_table.columns)
                        header_cell = header_table.cell(row_index, col_index)
                        header_cell.text = data
                        for paragraph in header_cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    print(
                        "The number of cells in the table does not match the length of headerList.")

            g_lastTable.createLastTableEmpty(docNumber)  # empty
            file_name = f'main{i + 1}.docx'
            document.save(file_name)
            file_names.append(file_name)
            break
        else:
            # Process normal chunks with 7 elements
            if len(document.tables) > 1:
                table = document.tables[1]
                num_rows = len(table.rows)
                if num_rows >= len(chunk):
                    for j in range(len(chunk)):
                        for k in range(len(chunk[j])):
                            cell = table.cell(j, k)
                            cell.margin_top = Pt(0)
                            cell.margin_right = Pt(0)
                            cell.text = chunk[j][k]
                            # Set alignment to right for the cell's paragraphs
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    print(
                        f"Table 1 in 'pt1.docx' does not have enough rows for chunk {i + 1}.")

        if len(document.tables) > 0:
            header_table = document.tables[0]
            if len(header_table.rows) * len(header_table.columns) == len(headerList):
                for j, data in enumerate(headerList):
                    row_index = j // len(header_table.columns)
                    col_index = j % len(header_table.columns)
                    header_cell = header_table.cell(row_index, col_index)
                    header_cell.text = data
                    for paragraph in header_cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                print(
                    "The number of cells in the table does not match the length of headerList.")

        file_name = f'main{i + 1}.docx'
        document.save(file_name)
        file_names.append(file_name)
    file_names.append('lastPage.docx')
    return file_names  # List of File Names
